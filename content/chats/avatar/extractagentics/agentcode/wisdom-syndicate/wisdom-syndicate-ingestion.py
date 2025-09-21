"""
Wisdom Syndicate Ingestion Agent
A comprehensive document ingestion and processing module for extracting structured data
from unstructured documents in the extractives industry.
"""

import re
import uuid
import logging
from datetime import datetime
from typing import List, Dict, Optional, Union, Tuple, Any
from dataclasses import dataclass, asdict
from enum import Enum
from pathlib import Path
import json

# External dependencies
import spacy
from spacy.matcher import Matcher
import pdfplumber
from docx import Document
from dateutil import parser as date_parser
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.chunk import ne_chunk
from nltk.tag import pos_tag

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
    nltk.download('maxent_ne_chunker', quiet=True)
    nltk.download('words', quiet=True)
except:
    pass

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# =================================
# Data Class Definitions
# =================================

class ProjectType(Enum):
    MINING = "mining"
    OIL_GAS = "oil_gas"
    INFRASTRUCTURE = "infrastructure"
    RENEWABLE_ENERGY = "renewable_energy"
    UNKNOWN = "unknown"

class IssueCategory(Enum):
    RESETTLEMENT = "resettlement"
    ENVIRONMENTAL = "environmental"
    COMMUNITY_ENGAGEMENT = "community_engagement"
    COMPENSATION = "compensation"
    CULTURAL_HERITAGE = "cultural_heritage"
    GRIEVANCE = "grievance"
    MONITORING = "monitoring"

class IssueSeverity(Enum):
    LOW = 1
    MEDIUM = 2
    HIGH = 3
    CRITICAL = 4

class StandardsFramework(Enum):
    IFC_PS = "IFC Performance Standards"
    WORLD_BANK_ESS = "World Bank ESS"
    EQUATOR_PRINCIPLES = "Equator Principles"
    NATIONAL_REGULATION = "National Regulation"
    INDUSTRY_STANDARD = "Industry Standard"
    OTHER = "Other"

@dataclass
class ProjectContext:
    project_id: str
    project_type: ProjectType
    location: Dict[str, str]  # {country, region, site}
    scale: str  # small, medium, large, mega
    timeline: Dict[str, Optional[datetime]]  # {start, expected_completion, actual_completion}
    stakeholders: List[str]
    regulatory_framework: List[StandardsFramework]
    environmental_context: Optional[str] = None
    social_context: Optional[str] = None
    economic_context: Optional[str] = None

@dataclass
class IssueRecord:
    issue_id: str
    project_context: ProjectContext
    category: IssueCategory
    severity: IssueSeverity
    description: str
    root_causes: List[str]
    timeline: Dict[str, Optional[datetime]]  # {identified, resolved}
    stakeholders_affected: List[str]
    resolution_approach: Optional[str] = None
    outcome: Optional[str] = None
    lessons_learned: List[str] = None
    cost_impact: Optional[float] = None
    time_impact: Optional[int] = None  # days
    recurrence_risk: Optional[float] = None  # 0-1 probability
    tags: List[str] = None

@dataclass
class BestPractice:
    practice_id: str
    category: IssueCategory
    title: str
    description: str
    implementation_steps: List[str]
    success_factors: List[str]
    applicable_contexts: List[str]
    evidence_base: List[str]  # References to projects/issues
    effectiveness_score: float  # 0-1

@dataclass
class StandardsEvolution:
    standard: StandardsFramework
    version: str
    effective_date: datetime
    key_changes: List[str]
    impact_areas: List[IssueCategory]
    adoption_rate: float  # 0-1

@dataclass
class PeerConnection:
    source_project: str
    target_project: str
    similarity_score: float  # 0-1
    shared_issues: List[IssueCategory]
    applicable_learnings: List[str]

# =================================
# NLP Patterns and Keywords
# =================================

EXTRACTIVE_KEYWORDS = {
    'project_type': {
        'mining': ['mining', 'mine', 'mineral', 'extraction', 'ore', 'copper', 'gold', 'coal', 'iron'],
        'oil_gas': ['oil', 'gas', 'petroleum', 'pipeline', 'refinery', 'drilling', 'offshore'],
        'infrastructure': ['infrastructure', 'road', 'railway', 'dam', 'power plant', 'transmission'],
        'renewable_energy': ['solar', 'wind', 'hydro', 'renewable', 'green energy']
    },
    'issue_category': {
        'resettlement': ['resettlement', 'relocation', 'displacement', 'land acquisition', 'eviction'],
        'environmental': ['environmental', 'pollution', 'contamination', 'biodiversity', 'deforestation', 'emissions'],
        'community_engagement': ['community engagement', 'consultation', 'participation', 'stakeholder', 'dialogue'],
        'compensation': ['compensation', 'payment', 'indemnity', 'reimbursement', 'livelihood restoration'],
        'cultural_heritage': ['cultural heritage', 'sacred site', 'archaeological', 'indigenous', 'traditional'],
        'grievance': ['grievance', 'complaint', 'dispute', 'conflict', 'protest', 'demonstration'],
        'monitoring': ['monitoring', 'evaluation', 'assessment', 'audit', 'compliance', 'inspection']
    },
    'severity_indicators': {
        'critical': ['critical', 'severe', 'emergency', 'crisis', 'urgent', 'catastrophic'],
        'high': ['high', 'significant', 'major', 'serious', 'substantial'],
        'medium': ['medium', 'moderate', 'average', 'normal'],
        'low': ['low', 'minor', 'minimal', 'slight', 'negligible']
    },
    'standards': {
        'IFC_PS': ['IFC', 'Performance Standard', 'PS1', 'PS2', 'PS3', 'PS4', 'PS5', 'PS6', 'PS7', 'PS8'],
        'WORLD_BANK_ESS': ['World Bank', 'ESS', 'Environmental and Social Standard', 'ESF'],
        'EQUATOR_PRINCIPLES': ['Equator Principles', 'EP', 'EPFI'],
        'NATIONAL_REGULATION': ['national', 'regulation', 'law', 'legislation', 'regulatory', 'government']
    },
    'stakeholders': [
        'local community', 'indigenous people', 'government', 'NGO', 'civil society',
        'contractor', 'supplier', 'employee', 'worker', 'union', 'media',
        'investor', 'shareholder', 'regulator', 'local authority', 'traditional leader'
    ]
}

# =================================
# Ingestion Agent Class
# =================================

class IngestionAgent:
    """
    Automated document ingestion and information extraction agent for the Wisdom Syndicate system.
    """
    
    def __init__(self, nlp_model: str = "en_core_web_sm"):
        """
        Initialize the Ingestion Agent.
        
        Args:
            nlp_model: spaCy model to use for NLP processing
        """
        try:
            self.nlp = spacy.load(nlp_model)
        except OSError:
            logger.warning(f"Model {nlp_model} not found. Downloading...")
            import subprocess
            subprocess.run(["python", "-m", "spacy", "download", nlp_model])
            self.nlp = spacy.load(nlp_model)
        
        # Add custom entity patterns
        self._setup_custom_patterns()
        
        logger.info("IngestionAgent initialized successfully")
    
    def _setup_custom_patterns(self):
        """Setup custom patterns for domain-specific entity recognition."""
        self.matcher = Matcher(self.nlp.vocab)
        
        # Project name patterns
        project_patterns = [
            [{"TEXT": {"REGEX": r"^[A-Z][a-z]+$"}}, {"TEXT": {"IN": ["Project", "Mine", "Field", "Site"]}}],
            [{"TEXT": "Project"}, {"IS_TITLE": True}]
        ]
        self.matcher.add("PROJECT_NAME", project_patterns)
        
        # Standards patterns
        standards_patterns = [
            [{"TEXT": {"IN": ["IFC", "World", "Equator"]}}, {"TEXT": {"IN": ["Performance", "Bank", "Principles"]}}],
            [{"TEXT": "PS"}, {"TEXT": {"REGEX": r"^\d$"}}],
            [{"TEXT": "ESS"}, {"TEXT": {"REGEX": r"^\d+$"}}]
        ]
        self.matcher.add("STANDARDS", standards_patterns)
    
    def ingest_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Main entry point for document ingestion.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted ProjectContext and IssueRecord instances
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Starting ingestion of: {file_path.name}")
        
        # Extract text based on file type
        if file_path.suffix.lower() == '.docx':
            text = self._extract_text_from_docx(file_path)
        elif file_path.suffix.lower() == '.pdf':
            text = self._extract_text_from_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        # Process text with NLP
        doc = self.nlp(text)
        
        # Extract structured information
        extracted_info = self._extract_information(doc, text)
        
        # Create data objects
        project_contexts = self._create_project_contexts(extracted_info)
        issue_records = self._create_issue_records(extracted_info, project_contexts)
        
        # Convert to dictionaries
        result = {
            'project_contexts': [asdict(pc) for pc in project_contexts],
            'issue_records': [asdict(ir) for ir in issue_records],
            'metadata': {
                'source_file': str(file_path),
                'extraction_date': datetime.now().isoformat(),
                'text_length': len(text),
                'entities_found': len(extracted_info.get('entities', []))
            }
        }
        
        logger.info(f"Successfully processed {file_path.name}: {len(project_contexts)} projects, {len(issue_records)} issues")
        
        return result
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            text_parts = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = '\n'.join([' | '.join(str(cell) for cell in row if cell) for row in table])
                            text_parts.append(table_text)
            
            return '\n\n'.join(text_parts)
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise
    
    def _extract_information(self, doc: spacy.tokens.Doc, raw_text: str) -> Dict[str, Any]:
        """Extract structured information from the processed document."""
        info = {
            'entities': [],
            'dates': [],
            'locations': [],
            'organizations': [],
            'stakeholders': [],
            'standards': [],
            'issues': [],
            'keywords': [],
            'sentences': []
        }
        
        # Extract entities
        for ent in doc.ents:
            info['entities'].append({
                'text': ent.text,
                'label': ent.label_,
                'start': ent.start_char,
                'end': ent.end_char
            })
            
            if ent.label_ in ['DATE']:
                info['dates'].append(ent.text)
            elif ent.label_ in ['GPE', 'LOC']:
                info['locations'].append(ent.text)
            elif ent.label_ in ['ORG']:
                info['organizations'].append(ent.text)
        
        # Extract custom patterns
        matches = self.matcher(doc)
        for match_id, start, end in matches:
            span = doc[start:end]
            match_label = self.nlp.vocab.strings[match_id]
            
            if match_label == "PROJECT_NAME":
                info['entities'].append({
                    'text': span.text,
                    'label': 'PROJECT',
                    'start': span.start_char,
                    'end': span.end_char
                })
            elif match_label == "STANDARDS":
                info['standards'].append(span.text)
        
        # Extract keywords and categorize issues
        sentences = sent_tokenize(raw_text)
        info['sentences'] = sentences
        
        for sent in sentences:
            sent_lower = sent.lower()
            
            # Check for stakeholders
            for stakeholder in EXTRACTIVE_KEYWORDS['stakeholders']:
                if stakeholder in sent_lower:
                    info['stakeholders'].append(stakeholder)
            
            # Check for issue categories
            for category, keywords in EXTRACTIVE_KEYWORDS['issue_category'].items():
                for keyword in keywords:
                    if keyword in sent_lower:
                        info['issues'].append({
                            'category': category,
                            'sentence': sent,
                            'keyword': keyword
                        })
                        info['keywords'].append(keyword)
            
            # Check for standards
            for standard_type, keywords in EXTRACTIVE_KEYWORDS['standards'].items():
                for keyword in keywords:
                    if keyword in sent:
                        info['standards'].append(standard_type)
        
        # Deduplicate
        info['stakeholders'] = list(set(info['stakeholders']))
        info['standards'] = list(set(info['standards']))
        info['keywords'] = list(set(info['keywords']))
        
        return info
    
    def _create_project_contexts(self, extracted_info: Dict[str, Any]) -> List[ProjectContext]:
        """Create ProjectContext objects from extracted information."""
        projects = []
        
        # Try to identify distinct projects
        project_names = [ent['text'] for ent in extracted_info['entities'] if ent['label'] == 'PROJECT']
        
        # If no projects identified, create a default one
        if not project_names:
            project_names = ['Default Project']
        
        for project_name in set(project_names):
            project_id = str(uuid.uuid4())
            
            # Determine project type
            project_type = ProjectType.UNKNOWN
            for ptype, keywords in EXTRACTIVE_KEYWORDS['project_type'].items():
                if any(kw in ' '.join(extracted_info['keywords']) for kw in keywords):
                    project_type = ProjectType(ptype)
                    break
            
            # Extract location
            location = {
                'country': extracted_info['locations'][0] if extracted_info['locations'] else 'Unknown',
                'region': extracted_info['locations'][1] if len(extracted_info['locations']) > 1 else 'Unknown',
                'site': project_name
            }
            
            # Parse dates for timeline
            timeline = {
                'start': None,
                'expected_completion': None,
                'actual_completion': None
            }
            
            for date_str in extracted_info['dates'][:3]:  # Take first 3 dates
                try:
                    parsed_date = date_parser.parse(date_str, fuzzy=True)
                    if timeline['start'] is None:
                        timeline['start'] = parsed_date
                    elif timeline['expected_completion'] is None:
                        timeline['expected_completion'] = parsed_date
                    else:
                        timeline['actual_completion'] = parsed_date
                except:
                    continue
            
            # Map standards
            regulatory_framework = []
            for std in extracted_info['standards']:
                if std in ['IFC_PS', 'WORLD_BANK_ESS', 'EQUATOR_PRINCIPLES', 'NATIONAL_REGULATION']:
                    regulatory_framework.append(StandardsFramework[std])
                else:
                    regulatory_framework.append(StandardsFramework.OTHER)
            
            # Create context snippets
            context_sentences = extracted_info['sentences'][:10]  # First 10 sentences
            
            project = ProjectContext(
                project_id=project_id,
                project_type=project_type,
                location=location,
                scale='medium',  # Default
                timeline=timeline,
                stakeholders=extracted_info['stakeholders'],
                regulatory_framework=regulatory_framework,
                environmental_context=' '.join([s for s in context_sentences if any(kw in s.lower() for kw in ['environment', 'pollution', 'biodiversity'])]),
                social_context=' '.join([s for s in context_sentences if any(kw in s.lower() for kw in ['community', 'social', 'stakeholder'])]),
                economic_context=' '.join([s for s in context_sentences if any(kw in s.lower() for kw in ['economic', 'cost', 'investment', 'budget'])])
            )
            
            projects.append(project)
        
        return projects
    
    def _create_issue_records(self, extracted_info: Dict[str, Any], project_contexts: List[ProjectContext]) -> List[IssueRecord]:
        """Create IssueRecord objects from extracted information."""
        issue_records = []
        
        # Group issues by category
        issues_by_category = {}
        for issue_info in extracted_info['issues']:
            category = issue_info['category']
            if category not in issues_by_category:
                issues_by_category[category] = []
            issues_by_category[category].append(issue_info)
        
        # Create issue records for each category
        for category, issues in issues_by_category.items():
            issue_id = str(uuid.uuid4())
            
            # Determine severity
            severity = IssueSeverity.MEDIUM  # Default
            severity_text = ' '.join([issue['sentence'] for issue in issues]).lower()
            
            for sev_level, indicators in EXTRACTIVE_KEYWORDS['severity_indicators'].items():
                if any(ind in severity_text for ind in indicators):
                    severity_map = {'critical': 4, 'high': 3, 'medium': 2, 'low': 1}
                    severity = IssueSeverity(severity_map[sev_level])
                    break
            
            # Extract description
            description = issues[0]['sentence'] if issues else f"{category} issue identified"
            
            # Extract dates for timeline
            timeline = {'identified': None, 'resolved': None}
            for date_str in extracted_info['dates']:
                try:
                    parsed_date = date_parser.parse(date_str, fuzzy=True)
                    if timeline['identified'] is None:
                        timeline['identified'] = parsed_date
                    elif timeline['resolved'] is None:
                        timeline['resolved'] = parsed_date
                    break
                except:
                    continue
            
            # Create issue record
            issue = IssueRecord(
                issue_id=issue_id,
                project_context=project_contexts[0] if project_contexts else None,  # Link to first project
                category=IssueCategory(category),
                severity=severity,
                description=description,
                root_causes=[issue['keyword'] for issue in issues[:3]],  # Top 3 keywords as causes
                timeline=timeline,
                stakeholders_affected=[s for s in extracted_info['stakeholders'] if any(s in issue['sentence'] for issue in issues)],
                resolution_approach=self._extract_resolution(issues, extracted_info['sentences']),
                outcome=self._extract_outcome(issues, extracted_info['sentences']),
                lessons_learned=self._extract_lessons(issues, extracted_info['sentences']),
                tags=[issue['keyword'] for issue in issues]
            )
            
            issue_records.append(issue)
        
        return issue_records
    
    def _extract_resolution(self, issues: List[Dict], sentences: List[str]) -> Optional[str]:
        """Extract resolution approach from surrounding context."""
        resolution_keywords = ['resolved', 'solution', 'addressed', 'implemented', 'mitigation', 'response']
        
        for sent in sentences:
            if any(kw in sent.lower() for kw in resolution_keywords):
                return sent
        
        return None
    
    def _extract_outcome(self, issues: List[Dict], sentences: List[str]) -> Optional[str]:
        """Extract outcome from surrounding context."""
        outcome_keywords = ['result', 'outcome', 'impact', 'effect', 'consequence', 'led to', 'resulted in']
        
        for sent in sentences:
            if any(kw in sent.lower() for kw in outcome_keywords):
                return sent
        
        return None
    
    def _extract_lessons(self, issues: List[Dict], sentences: List[str]) -> List[str]:
        """Extract lessons learned from surrounding context."""
        lesson_keywords = ['lesson', 'learned', 'recommendation', 'suggest', 'should', 'must', 'important to']
        lessons = []
        
        for sent in sentences:
            if any(kw in sent.lower() for kw in lesson_keywords):
                lessons.append(sent)
                if len(lessons) >= 3:  # Limit to 3 lessons
                    break
        
        return lessons

# =================================
# Example Usage
# =================================

def example_usage():
    """Demonstrate the usage of IngestionAgent."""
    
    # Initialize the agent
    agent = IngestionAgent()
    
    # Example 1: Process a DOCX file
    try:
        # Assuming we have a sample file
        sample_file = Path("120127_Benchmarking Report Antamina.docx")
        
        if sample_file.exists():
            results = agent.ingest_document(sample_file)
            
            print("=== Extraction Results ===")
            print(f"Projects found: {len(results['project_contexts'])}")
            print(f"Issues found: {len(results['issue_records'])}")
            
            # Display first project
            if results['project_contexts']:
                print("\n=== First Project ===")
                print(json.dumps(results['project_contexts'][0], indent=2, default=str))
            
            # Display first issue
            if results['issue_records']:
                print("\n=== First Issue ===")
                print(json.dumps(results['issue_records'][0], indent=2, default=str))
            
            # Save results to file
            output_file = Path("extracted_data.json")
            with open(output_file, 'w') as f:
                json.dump(results, f, indent=2, default=str)
            print(f"\nResults saved to: {output_file}")
        
        else:
            # Create dummy content for demonstration
            print("Sample file not found. Creating dummy content...")
            
            dummy_text = """
            Antamina Mining Project - Resettlement Impact Assessment Report
            
            Date: January 27, 2012
            Location: Ancash Region, Peru
            
            Executive Summary:
            The Antamina copper and zinc mine has implemented a comprehensive resettlement program 
            following IFC Performance Standard 5. The project affected approximately 200 families 
            in the local community who required relocation due to mine expansion.
            
            Critical Issue: Land acquisition resulted in significant community protests in March 2011.
            The company responded by enhancing compensation packages and implementing a grievance mechanism.
            
            Resolution: Through extensive community engagement and consultation, an agreement was reached
            in August 2011. The compensation included cash payments, livelihood restoration programs,
            and construction of new housing.
            
            Outcome: By December 2011, 180 families had successfully relocated with improved living
            conditions. The remaining 20 families are in ongoing negotiations.
            
            Lessons Learned:
            1. Early and continuous stakeholder engagement is essential for successful resettlement.
            2. Compensation must consider both economic and social impacts on affected communities.
            3. Grievance mechanisms should be established before project implementation begins.
            
            Environmental Context: The project area includes sensitive biodiversity zones requiring
            careful management and monitoring.
            
            Regulatory Framework: The project complies with IFC Performance Standards, World Bank ESS,
            and Peruvian national regulations on mining and resettlement.
            """
            
            # Save dummy content to a temporary file
            temp_file = Path("temp_dummy_report.txt")
            with open(temp_file, 'w') as f:
                f.write(dummy_text)
            
            # Process as text directly (simplified for demo)
            doc = agent.nlp(dummy_text)
            extracted_info = agent._extract_information(doc, dummy_text)
            project_contexts = agent._create_project_contexts(extracted_info)
            issue_records = agent._create_issue_records(extracted_info, project_contexts)
            
            print("\n=== Extracted from Dummy Content ===")
            print(f"Projects: {len(project_contexts)}")
            print(f"Issues: {len(issue_records)}")
            
            if project_contexts:
                print(f"\nProject Type: {project_contexts[0].project_type.value}")
                print(f"Location: {project_contexts[0].location}")
                print(f"Stakeholders: {project_contexts[0].stakeholders}")
            
            if issue_records:
                print(f"\nIssue Category: {issue_records[0].category.value}")
                print(f"Severity: {issue_records[0].severity.value}")
                print(f"Description: {issue_records[0].description[:100]}...")
            
            # Clean up
            temp_file.unlink()
    
    except Exception as e:
        logger.error(f"Error in example usage: {e}")
        raise

if __name__ == "__main__":
    example_usage()
